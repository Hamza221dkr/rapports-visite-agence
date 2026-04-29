"""
Moteur de génération — respecte fidèlement le modèle SYNTHESE COMPTE RENDU DE VISITE AGENCE.
Structure : paragraphes avec tirets (pas de tableaux pour les sections texte).
"""

import io
from copy import deepcopy
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Index colonnes Excel ────────────────────────────────────────────────────
C_THEME=0; C_QUEST=1; C_SCORE=2; C_NOTE=3; C_OBS=4; C_RESP=5
C_NTHEME=6; C_NGLOB=7; C_NOTAT=8

THEME_MAP = {
    "FONCTIONNEMENT": "Fonctionnement agence",
    "QUALITE":        "Qualité de service",
    "ENVIRONNEMENT":  "Sécurité & Risques opérationnels",
    "ORGANISATION":   "Organisation & Animation commerciale",
    "PERFORMANCES":   "Performance Commerciale",
    "MANAGEMENT":     "Management & Leadership",
}

# Libellés exacts du tableau du modèle
THEME_TABLE_LABELS = [
    ("Fonctionnement agence",       "FONCTIONNEMENT"),
    ("Qualité de service",          "QUALITE"),
    ("Sécurité & Risques opérationnels", "ENVIRONNEMENT"),
    ("Organisation du travail",     "ORGANISATION"),
    ("Performance Commerciale",     "PERFORMANCES"),
    ("Management & Leadership",     "MANAGEMENT"),
]

SUPPORT_KEYS = [
    ("DR-RMG","Direction Régionale / Moyens Généraux"),("DR","Direction Régionale"),
    ("MARKETING","Marketing"),("HAROUNA","Sécurité"),("SECURITE","Sécurité"),
    ("DSA","Direction Supports & Animation"),("RMG","Moyens Généraux"),
    ("INFORMATIQUE","Informatique"),("RMSQ","Qualité de Service"),
    ("CA","Chef Agence"),("CSO","Chef Agence"),("ROP","Chef Agence"),("CC","Chef Agence"),
]

# Couleurs pour la table de scoring uniquement
SCORE_COLORS = {
    "Bon":        {"bg":"C6EFCE","fg":"1E5631"},
    "Acceptable": {"bg":"FFEB9C","fg":"7D4E00"},
    "Critique":   {"bg":"FFC7CE","fg":"9C0006"},
    "—":          {"bg":"F2F2F2","fg":"595959"},
}

# ── Helpers données ─────────────────────────────────────────────────────────

def score_label(n):
    if n is None: return "—"
    return "Bon" if n >= 2.5 else ("Acceptable" if n >= 1.5 else "Critique")

def resolve_theme(raw: str) -> str:
    up = raw.upper()
    for k in THEME_MAP:
        if k in up: return k
    return raw

def resolve_entity(resp_raw: str) -> str:
    up = resp_raw.upper()
    for key, label in SUPPORT_KEYS:
        if key in up: return label
    return "Chef Agence"

# ── Helpers Word ────────────────────────────────────────────────────────────

def _shd(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _p_after(ref_para):
    """Insère un nouveau paragraphe juste après ref_para dans le doc XML."""
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    return new_p

def _add_para_after(doc, ref_para, text, bold=False, size=11,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, color_hex=None, italic=False):
    """Ajoute un paragraphe juste après ref_para en respectant le style du modèle."""
    new_p_xml = _p_after(ref_para)
    # Créer un objet paragraph python-docx autour du XML
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p_xml, ref_para._parent)
    new_para.alignment = align
    run = new_para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return new_para

def add_heading(doc, para, text):
    """Remplace le texte d'un paragraphe titre (bold)."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        para.runs[0].bold = True
    else:
        r = para.add_run(text)
        r.bold = True

# ── Lecture Excel ───────────────────────────────────────────────────────────

def list_agencies(file_bytes: bytes) -> list:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    return [n for n in wb.sheetnames if wb[n].max_row > 1]

def read_agency(ws) -> tuple:
    thematiques, current, note_globale, notation = {}, None, None, None
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[C_THEME]:
            current = str(row[C_THEME]).replace("\n", " ").strip()
        if not current: continue
        if current not in thematiques:
            thematiques[current] = {"note": None, "questions": []}
        if row[C_NTHEME] is not None:
            thematiques[current]["note"] = row[C_NTHEME]
        if row[C_NGLOB] and not note_globale:
            note_globale = row[C_NGLOB]
        if row[C_NOTAT] and not notation:
            notation = row[C_NOTAT]
        if row[C_QUEST]:
            thematiques[current]["questions"].append({
                "question":    str(row[C_QUEST]).replace("\n", " ").strip(),
                "scoring":     row[C_SCORE],
                "note":        row[C_NOTE],
                "obs":         str(row[C_OBS]).strip() if row[C_OBS] else "",
                "responsable": str(row[C_RESP]).strip() if row[C_RESP] else "",
            })
    return thematiques, note_globale, notation

# ── Construction du rapport individuel (fidèle au modèle) ──────────────────

def build_document(agency_name, thematiques, note_globale, notation_globale,
                   rz_name="", date_visite="", gpt_insights=None,
                   template_path=None) -> Document:

    # Charger le modèle s'il est disponible, sinon créer from scratch
    if template_path:
        try:
            doc = Document(template_path)
            return _fill_template(doc, agency_name, thematiques, note_globale,
                                  notation_globale, rz_name, date_visite, gpt_insights)
        except Exception:
            pass

    return _build_from_scratch(agency_name, thematiques, note_globale,
                               notation_globale, rz_name, date_visite, gpt_insights)


def _fill_template(doc, agency_name, thematiques, note_globale, notation_globale,
                   rz_name, date_visite, gpt_insights):
    """Remplit le modèle Word existant avec les données de l'agence."""
    paras = doc.paragraphs

    # ── En-tête ─────────────────────────────────────────────────────────────
    for p in paras:
        t = p.text.strip()
        if t.startswith("Agence"):
            _replace_para(p, "Agence : ", agency_name.upper())
        elif t.startswith("Date de visite"):
            _replace_para(p, "Date de visite : ", date_visite or "_______________")
        elif t.startswith("Responsable Agence"):
            _replace_para(p, "Responsable Agence : ", "_______________")
        elif t.startswith("Responsable de Zone"):
            _replace_para(p, "Responsable de Zone : ", rz_name or "_______________")

    # ── Score global & Niveau ───────────────────────────────────────────────
    ng_str  = f"{note_globale:.2f}" if note_globale else "—"
    notation = notation_globale or score_label(note_globale)

    for p in paras:
        if "Score global" in p.text:
            _replace_para(p, "Score global : ", f"{ng_str}/3", bold=True)
        elif p.text.strip().startswith("Niveau"):
            # Mettre en gras le niveau actuel
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = f"Niveau : {notation}"
                p.runs[0].bold = True

    # ── Table de scoring ─────────────────────────────────────────────────────
    if doc.tables:
        _fill_scoring_table(doc.tables[0], thematiques)

    # ── Collecter les données ─────────────────────────────────────────────────
    forts, alertes, actions_by_entity = _extract_data(thematiques)

    # ── Points forts ─────────────────────────────────────────────────────────
    pf_anchor = _find_para(paras, "Points forts")
    if pf_anchor:
        _remove_placeholder_bullets(paras, pf_anchor)
        _insert_bullets(doc, pf_anchor, forts)

    # ── Points d'alerte ──────────────────────────────────────────────────────
    pa_anchor = _find_para(paras, "alerte")
    if pa_anchor:
        _remove_placeholder_bullets(paras, pa_anchor)
        _insert_bullets(doc, pa_anchor, alertes)

    # ── Plan d'actions ────────────────────────────────────────────────────────
    plan_anchor = _find_para(paras, "actions prioritaire")
    if plan_anchor:
        _remove_placeholder_bullets(paras, plan_anchor)
        _insert_plan_bullets(doc, plan_anchor, actions_by_entity)

    # ── Appréciation générale ─────────────────────────────────────────────────
    appr_anchor = _find_para(paras, "Appréciation générale")
    if appr_anchor:
        conclusion = gpt_insights or _auto_conclusion(
            agency_name, thematiques, note_globale, notation_globale)
        _replace_appr(appr_anchor, conclusion)

    return doc


def _build_from_scratch(agency_name, thematiques, note_globale, notation_globale,
                        rz_name, date_visite, gpt_insights):
    """Construit le rapport en respectant la structure exacte du modèle."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.5)

    def p(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          color=None, space_before=0, space_after=4):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = para.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        return para

    ng_str  = f"{note_globale:.2f}" if note_globale else "—"
    notation = notation_globale or score_label(note_globale)

    # ── EN-TÊTE ──────────────────────────────────────────────────────────────
    p(f"Agence : {agency_name.upper()}", bold=True, size=13,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    p(f"Date de visite : {date_visite or '_______________'}")
    p("Responsable Agence : _______________")
    p(f"Responsable de Zone : {rz_name or '_______________'}")
    p()

    # ── SCORING GLOBAL ────────────────────────────────────────────────────────
    p("SCORING GLOBAL", bold=True, size=12, space_before=6, space_after=4)
    p(f"Score global : {ng_str}/3", bold=True)
    p()

    # Table de scoring
    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    # En-tête
    for ci, h in enumerate(["Critères", "Appréciation", "Note /3"]):
        cell = t.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shd(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
    # Données
    _fill_scoring_table_rows(t, thematiques, start_row=1)
    p()

    p(f"Niveau : {notation}", bold=True, space_after=8)

    # ── POINTS FORTS ──────────────────────────────────────────────────────────
    p("Points forts", bold=True, size=12, space_before=6, space_after=3)
    forts, alertes, actions_by_entity = _extract_data(thematiques)
    if forts:
        for ligne in forts:
            p(f"- {ligne}", size=11)
    else:
        p("- Aucun point fort identifié sur cette visite.", size=11)
    p()

    # ── POINTS D'ALERTE ───────────────────────────────────────────────────────
    p("Points d'alerte", bold=True, size=12, space_before=6, space_after=3)
    if alertes:
        for ligne in alertes:
            p(f"- {ligne}", size=11)
    else:
        p("- Aucun point d'alerte majeur identifié.", size=11)
    p()

    # ── PLAN D'ACTIONS PRIORITAIRE ────────────────────────────────────────────
    p("Plan d'actions prioritaire", bold=True, size=12, space_before=6, space_after=3)
    if actions_by_entity:
        for entity, actions in sorted(actions_by_entity.items()):
            p(f"{entity} :", bold=True, size=11, space_before=4, space_after=2)
            for action in actions:
                txt = f"- {action['action']}"
                if action['resp']:
                    txt += f"  (Resp. : {action['resp']})"
                if action['priorite'] == 'Urgente':
                    txt += "  ⚠ Urgent"
                p(txt, size=11)
    else:
        p("- Aucune action prioritaire définie.", size=11)
    p()

    # ── APPRÉCIATION GÉNÉRALE ─────────────────────────────────────────────────
    conclusion = gpt_insights or _auto_conclusion(
        agency_name, thematiques, note_globale, notation_globale)
    appr = doc.add_paragraph()
    appr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    appr.paragraph_format.space_before = Pt(8)
    r1 = appr.add_run("Appréciation générale : ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = appr.add_run(conclusion)
    r2.font.size = Pt(11)

    # Signature
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.paragraph_format.space_before = Pt(20)
    rs = sign.add_run(f"Le Responsable de Zone\n{rz_name or '_______________'}")
    rs.bold = True; rs.font.size = Pt(11)

    return doc


# ── Extraction des données ──────────────────────────────────────────────────

def _extract_data(thematiques):
    """Retourne (forts, alertes, actions_by_entity) — format texte tirets."""
    forts   = []
    alertes = []
    actions_by_entity = {}

    for theme_raw, td in thematiques.items():
        theme_label = THEME_MAP.get(resolve_theme(theme_raw), theme_raw)
        for q in td["questions"]:
            obs  = q["obs"] if q["obs"] and q["obs"] != "nan" else ""
            resp = q["responsable"]

            # Points forts
            if q["scoring"] == "Bon":
                detail = obs if obs and obs.upper() != "OK" else "Conforme aux standards"
                forts.append(f"[{theme_label}] {q['question'][:80]} → {detail}")

            # Points d'alerte
            elif q["scoring"] == "À améliorer":
                ligne = f"[{theme_label}] {q['question'][:80]}"
                if obs:
                    ligne += f" : {obs[:120]}"
                if resp:
                    ligne += f"  (Resp. : {resp})"
                alertes.append(ligne)

            # Plan d'actions
            if q["scoring"] in ("À améliorer", "Acceptable") and obs:
                entity = resolve_entity(resp)
                actions_by_entity.setdefault(entity, []).append({
                    "action":   obs,
                    "resp":     resp,
                    "priorite": "Urgente" if q["scoring"] == "À améliorer" else "Normale",
                    "theme":    theme_label,
                })

    # Trier : Urgentes en premier dans chaque entité
    for entity in actions_by_entity:
        actions_by_entity[entity].sort(key=lambda x: x["priorite"] == "Normale")

    return forts, alertes, actions_by_entity


# ── Remplissage de la table de scoring ─────────────────────────────────────

def _fill_scoring_table(table, thematiques):
    """Remplit le tableau de scoring du modèle (déjà chargé depuis le fichier)."""
    _fill_scoring_table_rows(table, thematiques, start_row=1)


def _fill_scoring_table_rows(table, thematiques, start_row=1):
    for ri, (label, key) in enumerate(THEME_TABLE_LABELS):
        row_idx = start_row + ri
        if row_idx >= len(table.rows):
            break
        row  = table.rows[row_idx]
        td   = next((v for k, v in thematiques.items() if key in k.upper()), None)
        nv   = td["note"] if td else None
        appr = score_label(nv)
        nstr = f"{nv:.2f}" if nv else "—"
        colors = SCORE_COLORS.get(appr, SCORE_COLORS["—"])

        # Critères (col 0)
        _set_cell_text(row.cells[0], label, bold=False)
        # Appréciation (col 1)
        _set_cell_text(row.cells[1], appr, bold=True, fg=colors["fg"],
                       bg=colors["bg"], center=True)
        # Note (col 2)
        _set_cell_text(row.cells[2], nstr, bold=True, fg=colors["fg"],
                       bg=colors["bg"], center=True)


def _set_cell_text(cell, text, bold=False, fg=None, bg=None, center=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if fg:
        run.font.color.rgb = RGBColor.from_string(fg)
    if bg:
        _shd(cell, bg)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Helpers pour remplissage du modèle ─────────────────────────────────────

def _find_para(paras, keyword):
    for p in paras:
        if keyword.lower() in p.text.lower():
            return p
    return None

def _replace_para(para, label, value, bold=False):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = label
        para.runs[0].bold = False
        r2 = para.add_run(value)
        r2.bold = bold
    else:
        r1 = para.add_run(label)
        r2 = para.add_run(value)
        r2.bold = bold

def _replace_appr(para, conclusion):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = "Appréciation générale : "
        para.runs[0].bold = True
        r2 = para.add_run(conclusion)
        r2.bold = False
    else:
        r1 = para.add_run("Appréciation générale : ")
        r1.bold = True
        r2 = para.add_run(conclusion)

def _remove_placeholder_bullets(paras, anchor):
    """Supprime les '- ' placeholder juste après anchor."""
    idx = list(paras).index(anchor)
    for p in list(paras)[idx+1:]:
        if p.text.strip() in ("- ", "-") and not any(r.bold for r in p.runs):
            p.clear()
        else:
            break

def _insert_bullets(doc, anchor, lines):
    """Insère des lignes '- texte' juste après anchor dans le doc."""
    from docx.text.paragraph import Paragraph
    last_para = anchor
    for line in lines:
        new_p_xml = _p_after(last_para)
        new_para = Paragraph(new_p_xml, anchor._parent)
        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = new_para.add_run(f"- {line}")
        run.font.size = Pt(11)
        last_para = new_para
    if not lines:
        new_p_xml = _p_after(last_para)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p_xml, anchor._parent)
        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_para.add_run("- Aucun élément identifié.").font.size = Pt(11)

def _insert_plan_bullets(doc, anchor, actions_by_entity):
    """Insère le plan d'actions groupé par entité support."""
    from docx.text.paragraph import Paragraph
    last_para = anchor
    if not actions_by_entity:
        new_p_xml = _p_after(last_para)
        np = Paragraph(new_p_xml, anchor._parent)
        np.add_run("- Aucune action prioritaire définie.").font.size = Pt(11)
        return
    for entity, actions in sorted(actions_by_entity.items()):
        # Sous-titre entité
        new_p_xml = _p_after(last_para)
        np = Paragraph(new_p_xml, anchor._parent)
        np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        np.paragraph_format.space_before = Pt(4)
        r = np.add_run(f"{entity} :")
        r.bold = True; r.font.size = Pt(11)
        last_para = np
        # Actions
        for a in actions:
            new_p_xml = _p_after(last_para)
            np2 = Paragraph(new_p_xml, anchor._parent)
            np2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            txt = f"- {a['action']}"
            if a['resp']:
                txt += f"  (Resp. : {a['resp']})"
            if a['priorite'] == 'Urgente':
                txt += "  ⚠ Urgent"
            r2 = np2.add_run(txt)
            r2.font.size = Pt(11)
            last_para = np2


# ── Conclusion automatique ──────────────────────────────────────────────────

def _auto_conclusion(agency_name, thematiques, note_globale, notation_globale):
    ng_str  = f"{note_globale:.2f}/3" if note_globale else "—"
    notation = notation_globale or score_label(note_globale)
    bons  = [THEME_MAP.get(resolve_theme(tr), tr) for tr, td in thematiques.items()
             if td.get("note") is not None and td["note"] >= 2.5]
    crits = [THEME_MAP.get(resolve_theme(tr), tr) for tr, td in thematiques.items()
             if td.get("note") is not None and td["note"] < 1.5]
    nb_a  = sum(1 for td in thematiques.values()
                for q in td["questions"] if q["scoring"] == "À améliorer")

    if notation == "Bon":
        appr_txt = "L'agence présente de bonnes pratiques et un niveau de performance satisfaisant. "
    elif notation == "Acceptable":
        appr_txt = "L'agence présente un niveau de fonctionnement acceptable avec des marges d'amélioration notables. "
    else:
        appr_txt = "Des insuffisances importantes ont été relevées, nécessitant des actions correctrices immédiates. "

    return (
        f"La visite de l'agence {agency_name} fait ressortir une note globale de {ng_str} ({notation}). "
        + appr_txt
        + (f"Points forts : {', '.join(bons)}. " if bons else "")
        + (f"Axes d'amélioration prioritaires : {', '.join(crits)}. " if crits
           else (f"{nb_a} point(s) d'alerte nécessitent un suivi rapproché. " if nb_a else ""))
        + "Le Chef Agence est invité à mettre en œuvre le plan d'actions défini, "
          "en liaison avec les entités support (Informatique, Moyens Généraux, Direction Régionale), "
          "avec un point de suivi mensuel."
    )


# ── Rapport consolidé ───────────────────────────────────────────────────────

def build_consolidated_report(all_agencies: list, rz_name="", date_visite="",
                               gpt_synthesis=None) -> Document:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(1.8)
        sec.left_margin = sec.right_margin = Cm(2.2)

    def p(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          color=None, space_before=0, space_after=4):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = para.add_run(text)
            run.bold = bold; run.font.size = Pt(size)
            if color: run.font.color.rgb = RGBColor.from_string(color)
        return para

    n = len(all_agencies)
    avg = sum(ag.get("note_globale") or 0 for ag in all_agencies) / max(n, 1)

    # ── TITRE ────────────────────────────────────────────────────────────────
    p("RAPPORT CONSOLIDÉ — COMPTES RENDUS DE VISITE DE ZONE", bold=True, size=14,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p(f"Direction Régionale — {n} agence(s) visitée(s)", size=11,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    p(f"Date de visite : {date_visite or '_______________'}     |     "
      f"Responsable de Zone : {rz_name or '_______________'}",
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # ── TABLEAU COMPARATIF DES SCORES ────────────────────────────────────────
    p("SCORING COMPARATIF DES AGENCES", bold=True, size=12, space_before=6, space_after=4)

    n_cols = n + 2  # Critères + Pondération + N agences
    ct = doc.add_table(rows=len(THEME_TABLE_LABELS) + 3, cols=n_cols)
    ct.style = "Table Grid"

    # En-tête
    _set_cell_text(ct.rows[0].cells[0], "Thématique", bold=True, fg="FFFFFF", bg="1F4E79", center=True)
    _set_cell_text(ct.rows[0].cells[1], "Poids", bold=True, fg="FFFFFF", bg="1F4E79", center=True)
    for ci, ag in enumerate(all_agencies, 2):
        _set_cell_text(ct.rows[0].cells[ci], ag["name"], bold=True, fg="FFFFFF", bg="1F4E79", center=True)

    # Lignes thématiques
    for ri, (label, key) in enumerate(THEME_TABLE_LABELS, 1):
        _set_cell_text(ct.rows[ri].cells[0], label, bg="EBF3FB")
        _set_cell_text(ct.rows[ri].cells[1], "1/6", center=True, bg="F2F2F2")
        for ci, ag in enumerate(all_agencies, 2):
            td   = next((v for k, v in ag["thematiques"].items() if key in k.upper()), None)
            nv   = td["note"] if td else None
            appr = score_label(nv)
            nstr = f"{nv:.2f}" if nv else "—"
            colors = SCORE_COLORS.get(appr, SCORE_COLORS["—"])
            _set_cell_text(ct.rows[ri].cells[ci], f"{nstr}  ({appr})",
                           bold=True, fg=colors["fg"], bg=colors["bg"], center=True)

    # Ligne note globale
    ri_g = len(THEME_TABLE_LABELS) + 1
    _set_cell_text(ct.rows[ri_g].cells[0], "NOTE GLOBALE", bold=True, fg="FFFFFF", bg="2E74B5")
    _set_cell_text(ct.rows[ri_g].cells[1], "", bg="2E74B5")
    for ci, ag in enumerate(all_agencies, 2):
        ng   = ag.get("note_globale")
        appr = ag.get("notation_globale") or score_label(ng)
        nstr = f"{ng:.2f}" if ng else "—"
        colors = SCORE_COLORS.get(appr, SCORE_COLORS["—"])
        _set_cell_text(ct.rows[ri_g].cells[ci], f"{nstr}  ({appr})",
                       bold=True, fg=colors["fg"], bg=colors["bg"], center=True)

    # Ligne classement
    ri_r = len(THEME_TABLE_LABELS) + 2
    _set_cell_text(ct.rows[ri_r].cells[0], "CLASSEMENT", bold=True, fg="FFFFFF", bg="1F4E79")
    _set_cell_text(ct.rows[ri_r].cells[1], "", bg="1F4E79")
    scores = sorted(all_agencies, key=lambda x: -(x.get("note_globale") or 0))
    ranks  = {ag["name"]: i+1 for i, ag in enumerate(scores)}
    medals = ["1er 🥇", "2e 🥈", "3e 🥉"] + [f"{i+1}e" for i in range(3, 20)]
    for ci, ag in enumerate(all_agencies, 2):
        rank = ranks.get(ag["name"], "-")
        label = medals[rank-1] if isinstance(rank, int) and rank <= len(medals) else str(rank)
        _set_cell_text(ct.rows[ri_r].cells[ci], label,
                       bold=True, fg="FFFFFF", bg="1F4E79", center=True)
    p()

    p(f"Note moyenne de zone : {avg:.2f}/3  ({score_label(avg)})",
      bold=True, space_after=8)

    # ── POINTS FORTS INTER-AGENCES ────────────────────────────────────────────
    p("Points forts", bold=True, size=12, space_before=6, space_after=3)
    forts_map = {}
    for ag in all_agencies:
        for tr, td in ag["thematiques"].items():
            label = THEME_MAP.get(resolve_theme(tr), tr)
            for q in td["questions"]:
                if q["scoring"] == "Bon":
                    key2 = (label, q["question"][:60])
                    forts_map.setdefault(key2, []).append(ag["name"])
    forts_sorted = sorted(forts_map.items(), key=lambda x: -len(x[1]))
    for (theme, quest), agencies in forts_sorted:
        txt = f"[{theme}] {quest[:80]}  — Agences : {' | '.join(agencies)}"
        p(f"- {txt}")
    if not forts_sorted:
        p("- Aucun point fort commun identifié.")
    p()

    # ── POINTS D'ALERTE INTER-AGENCES ─────────────────────────────────────────
    p("Points d'alerte", bold=True, size=12, space_before=6, space_after=3)
    alertes_map = {}
    for ag in all_agencies:
        for tr, td in ag["thematiques"].items():
            label = THEME_MAP.get(resolve_theme(tr), tr)
            for q in td["questions"]:
                if q["scoring"] == "À améliorer":
                    obs = q["obs"] if q["obs"] else ""
                    key3 = (label, q["question"][:60])
                    alertes_map.setdefault(key3, {"agencies": [], "obs": obs})
                    if ag["name"] not in alertes_map[key3]["agencies"]:
                        alertes_map[key3]["agencies"].append(ag["name"])
    alertes_sorted = sorted(alertes_map.items(), key=lambda x: -len(x[1]["agencies"]))
    for (theme, quest), info in alertes_sorted:
        nb = len(info["agencies"])
        recur = " ⚠ GÉNÉRALISÉ" if nb == n else (f" [{nb}/{n} agences]" if nb > 1 else "")
        txt = f"[{theme}] {quest[:80]}{recur}  — {' | '.join(info['agencies'])}"
        if info["obs"]:
            txt += f" : {info['obs'][:100]}"
        p(f"- {txt}")
    if not alertes_sorted:
        p("- Aucun point d'alerte commun identifié.")
    p()

    # ── PLAN D'ACTIONS CONSOLIDÉ ───────────────────────────────────────────────
    p("Plan d'actions prioritaire", bold=True, size=12, space_before=6, space_after=3)

    all_actions_map = {}
    for ag in all_agencies:
        for tr, td in ag["thematiques"].items():
            label = THEME_MAP.get(resolve_theme(tr), tr)
            for q in td["questions"]:
                if q["scoring"] in ("À améliorer", "Acceptable") and q["obs"]:
                    entity = resolve_entity(q["responsable"])
                    prio   = "Urgente" if q["scoring"] == "À améliorer" else "Normale"
                    akey   = (entity, q["obs"][:80])
                    rec    = all_actions_map.setdefault(akey, {
                        "entity": entity, "action": q["obs"],
                        "resp": q["responsable"], "priorite": prio, "agencies": []
                    })
                    if ag["name"] not in rec["agencies"]:
                        rec["agencies"].append(ag["name"])

    by_entity = {}
    for (entity, _), action in all_actions_map.items():
        by_entity.setdefault(entity, []).append(action)
    for ent in by_entity:
        by_entity[ent].sort(key=lambda x: x["priorite"] == "Normale")

    for entity, actions in sorted(by_entity.items()):
        p(f"{entity} :", bold=True, size=11, space_before=4, space_after=2)
        for a in actions:
            txt = f"- {a['action']}"
            if a['resp']:
                txt += f"  (Resp. : {a['resp']})"
            if a['priorite'] == 'Urgente':
                txt += "  ⚠ Urgent"
            txt += f"  [{' | '.join(a['agencies'])}]"
            p(txt)
    p()

    # ── APPRÉCIATION GÉNÉRALE ─────────────────────────────────────────────────
    conclusion = gpt_synthesis or _auto_conclusion_zone(all_agencies, avg, alertes_sorted)

    appr = doc.add_paragraph()
    appr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    appr.paragraph_format.space_before = Pt(8)
    r1 = appr.add_run("Appréciation générale : ")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = appr.add_run(conclusion)
    r2.font.size = Pt(11)

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.paragraph_format.space_before = Pt(20)
    rs = sign.add_run(f"Le Responsable de Zone\n{rz_name or '_______________'}")
    rs.bold = True; rs.font.size = Pt(11)

    return doc


def _auto_conclusion_zone(all_agencies, avg, alertes_sorted):
    best  = max(all_agencies, key=lambda x: x.get("note_globale") or 0)
    worst = min(all_agencies, key=lambda x: x.get("note_globale") or 0)
    recurrent = [quest for (theme, quest), info in alertes_sorted
                 if len(info["agencies"]) > 1][:2]
    return (
        f"La tournée de visite des {len(all_agencies)} agences de la zone fait ressortir une note moyenne "
        f"de {avg:.2f}/3 ({score_label(avg)}). "
        f"L'agence {best['name']} se distingue avec la meilleure performance "
        f"({best.get('note_globale', 0):.2f}/3), tandis que l'agence {worst['name']} nécessite "
        f"un accompagnement renforcé ({worst.get('note_globale', 0):.2f}/3). "
        + (f"Des points d'alerte récurrents ont été identifiés sur plusieurs agences : "
           f"{'; '.join(recurrent)}. " if recurrent else "")
        + "Un plan d'actions de zone est à mettre en place en concertation avec les entités support "
          "(Informatique, Moyens Généraux, Direction Régionale) pour traiter les problématiques "
          "transversales. Un suivi mensuel est recommandé."
    )


# ── Fonctions de génération principale ─────────────────────────────────────

def generate_reports(file_bytes: bytes, rz_name="", date_visite="",
                     agencies=None, gpt_fn=None, template_path=None) -> list:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    results = []
    for sheet_name in wb.sheetnames:
        if agencies and sheet_name.upper() not in [a.upper() for a in agencies]:
            continue
        ws = wb[sheet_name]
        if ws.max_row <= 1:
            continue
        thematiques, note_globale, notation_globale = read_agency(ws)

        gpt_insights = None
        if gpt_fn:
            try:
                gpt_insights = gpt_fn("individual", sheet_name, thematiques,
                                      note_globale, notation_globale)
            except Exception:
                pass

        doc = build_document(sheet_name, thematiques, note_globale, notation_globale,
                             rz_name, date_visite, gpt_insights, template_path)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        results.append((sheet_name, buf.read()))
    return results


def generate_consolidated(file_bytes: bytes, rz_name="", date_visite="",
                           agencies=None, gpt_fn=None) -> bytes:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_agencies = []
    for sheet_name in wb.sheetnames:
        if agencies and sheet_name.upper() not in [a.upper() for a in agencies]:
            continue
        ws = wb[sheet_name]
        if ws.max_row <= 1:
            continue
        thematiques, note_globale, notation_globale = read_agency(ws)
        all_agencies.append({"name": sheet_name, "thematiques": thematiques,
                              "note_globale": note_globale,
                              "notation_globale": notation_globale})

    gpt_synthesis = None
    if gpt_fn and all_agencies:
        try:
            gpt_synthesis = gpt_fn("consolidated", all_agencies)
        except Exception:
            pass

    doc = build_consolidated_report(all_agencies, rz_name, date_visite, gpt_synthesis)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
